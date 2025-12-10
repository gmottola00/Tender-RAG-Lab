"""DOCX parsing utilities using python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Union

from docx import Document  # type: ignore

from .normalizer import normalize_text

BlockDict = Dict[str, Any]


def _validate_docx_path(path: Union[str, Path]) -> Path:
    """Validate that the DOCX path exists and is a file."""
    docx_path = Path(path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found at {docx_path}")
    if not docx_path.is_file():
        raise ValueError(f"Path is not a file: {docx_path}")
    return docx_path


def parse_docx(path: Union[str, Path]) -> Dict[str, List[BlockDict]]:
    """Parse paragraphs and tables from a DOCX file.

    Args:
        path: Path to the DOCX file.

    Returns:
        Dictionary containing ``blocks`` (paragraphs/headings) and ``tables``.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the path is not a file.
        RuntimeError: If python-docx cannot open the document.
    """
    docx_path = _validate_docx_path(path)

    try:
        doc = Document(docx_path.as_posix())
    except Exception as exc:  # pragma: no cover - passthrough
        raise RuntimeError(f"Failed to open DOCX: {docx_path}") from exc

    blocks: List[BlockDict] = []
    for para in doc.paragraphs:
        text = normalize_text(para.text).strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        block_type = "paragraph"
        level = None

        if style_name.startswith("Heading"):
            block_type = "heading"
            try:
                level = int(style_name.replace("Heading", "").strip())
            except Exception:
                level = 1

        blocks.append(
            {
                "type": block_type,
                "level": level,
                "text": text,
                "style": style_name,
            }
        )

    tables: List[BlockDict] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [normalize_text(cell.text).strip() for cell in row.cells]
            rows.append(cells)
        text_rows = [" | ".join(r) for r in rows if any(r)]
        tables.append(
            {
                "type": "table_block",
                "raw_cells": rows,
                "text": "\n".join(text_rows).strip(),
            }
        )

    return {"blocks": blocks, "tables": tables}


__all__ = ["parse_docx"]
